import flask
import json
import wget
import datetime
import requests
from docx import Document
import os
from time import sleep
import convertapi

app = flask.Flask(__name__)

@app.route('/')
def index():
    return flask.render_template("index.html")

@app.route("/image", methods=['POST'])
def img():
    date = flask.request.form["date"]
    date_form = datetime.datetime.strptime(date, '%Y-%m-%d').date()
    date_today = datetime.date.today()
    if date_form > date_today:
        flask.abort(404)
    url = "https://api.nasa.gov/planetary/apod?date=" + date + "&hd=True&api_key=u6vr4iu4MnBKhpZO9o3HLyoMwQqKTS1TyIR5kxHi"
    img_url = requests.get(url).json().get('hdurl')
    caption = requests.get(url).json().get('explanation')
    wget.download(img_url, out="nasa.jpg")
    document = Document()
    document.add_picture("nasa.jpg")
    document.add_paragraph(caption)
    document.save('nasa.docx')
    os.remove("nasa.jpg")
#Use convert api
    convertapi.api_secret = 'yNNCmwvXeDwST4yd'
    convertapi.convert('pdf', {
    'File': 'nasa.docx'
    }, from_format = 'docx').save_files('nasa.pdf')
    return flask.render_template("image.html", src = img_url)
    os.remove("nasa.docx")
@app.route("/pdf", methods=['GET'])
def rr():
    return flask.send_file("nasa.pdf")
    os.remove("nasa.pdf")

if __name__ == '__main__':
	app.run()
